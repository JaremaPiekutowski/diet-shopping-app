from decimal import Decimal
from django.core.management.base import BaseCommand
from docx import Document
from planner.models import Plan
from recipes.models import IngredientCategory, RecipeIngredient


class Command(BaseCommand):
    help = "Exports plan and shopping list to docx files"

    def add_arguments(self, parser):
        parser.add_argument("plan_id", type=str, help="ID of the plan to be exported")

    def handle(self, *args, **kwargs):
        plan_id = kwargs["plan_id"]
        plan_to_export = Plan.objects.get(id=plan_id)
        if plan_to_export:
            self.do_export(plan_to_export)
        else:
            print("Plan not found")

    def get_plan_data(self, plan):
        plandays = plan.planday_set.all()
        plan_data = {
            "date_start": plan.date_start.date,
            "date_end": plan.date_end.date,
            "days": [
                {
                    "date": planday.day.date,
                    "meals": {
                        "breakfast": {
                            "title": (
                                planday.day.breakfast.title
                                if planday.day.breakfast
                                else "No breakfast planned"
                            ),
                            "ingredients": (
                                [
                                    ingredient
                                    for ingredient in planday.day.breakfast.ingredients
                                ]
                                if planday.day.breakfast
                                else ""
                            ),
                            "cooking_instructions": (
                                planday.day.breakfast.cooking_instructions
                                if planday.day.breakfast
                                else ""
                            ),
                        },
                        "lunch": {
                            "title": (
                                planday.day.lunch.title
                                if planday.day.lunch
                                else "No lunch planned"
                            ),
                            "ingredients": (
                                [
                                    ingredient
                                    for ingredient in planday.day.lunch.ingredients
                                ]
                                if planday.day.lunch
                                else ""
                            ),
                            "cooking_instructions": (
                                planday.day.lunch.cooking_instructions
                                if planday.day.lunch
                                else ""
                            ),
                        },
                        "dinner": {
                            "title": (
                                planday.day.dinner.title
                                if planday.day.dinner
                                else "No dinner planned"
                            ),
                            "ingredients": (
                                [
                                    ingredient
                                    for ingredient in planday.day.dinner.ingredients
                                ]
                                if planday.day.dinner
                                else ""
                            ),
                            "cooking_instructions": (
                                planday.day.dinner.cooking_instructions
                                if planday.day.dinner
                                else ""
                            ),
                        },
                        "supper": {
                            "title": (
                                planday.day.supper.title
                                if planday.day.supper
                                else "No supper planned"
                            ),
                            "ingredients": (
                                [
                                    ingredient
                                    for ingredient in planday.day.supper.ingredients
                                ]
                                if planday.day.supper
                                else ""
                            ),
                            "cooking_instructions": (
                                planday.day.supper.cooking_instructions
                                if planday.day.supper
                                else ""
                            ),
                        },
                    },
                }
                for planday in plandays
            ],
        }
        return plan_data

    def get_date_start_end(self, plan):
        # Use re to extract everything before "," from the date string
        plan_data = self.get_plan_data(plan)
        return plan_data["date_start"].strftime("%d.%m.%Y"), plan_data[
            "date_end"
        ].strftime("%d.%m.%Y")

    def get_shopping_list_data(self, plan):
        shopping_list = {}

        for category, _ in IngredientCategory.choices:
            shopping_list[category] = []

        plandays = plan.planday_set.all()

        for planday in plandays:
            day = planday.day

            meals = [day.breakfast, day.lunch, day.dinner, day.supper]

            for meal_recipe in meals:
                if meal_recipe:
                    recipe_ingredients = RecipeIngredient.objects.filter(
                        recipe=meal_recipe
                    )
                    for ri in recipe_ingredients:
                        ingredient = ri.ingredient

                        category = ingredient.category or IngredientCategory.OTHER

                        # Check if the ingredient is already in the shopping list
                        existing_ingredient = next(
                            (
                                item
                                for item in shopping_list[category]
                                if item["name"] == ingredient.name
                            ),
                            None,
                        )

                        if existing_ingredient:
                            existing_ingredient["total_grams"] += Decimal(
                                str(ri.quantity_grams)
                            )
                            existing_ingredient["total_price"] += Decimal(
                                str(ri.quantity_grams)
                            ) * (ingredient.price_per_hundred_gram / 100)
                        else:
                            shopping_list[category].append(
                                {
                                    "name": ingredient.name,
                                    "total_grams": Decimal(str(ri.quantity_grams)),
                                    "total_price": (
                                        (
                                            Decimal(str(ri.quantity_grams))
                                            * (ingredient.price_per_hundred_gram / 100)
                                        )
                                    ),
                                }
                            )
        return shopping_list

    def export_plan(self, plan):
        date_start, date_end = self.get_date_start_end(plan)
        plan_data = self.get_plan_data(plan)
        doc = Document()
        doc.add_heading(f"Diet plan for {date_start} - {date_end}", 0)
        for day in plan_data["days"]:
            doc.add_heading(day["date"].strftime("%d %B %Y, %A"), level=1)
            doc.add_paragraph()
            for meal, title in day["meals"].items():
                doc.add_paragraph(f"{meal.capitalize()}: {title}")
        filename = (
            f"exported_data/plan_{date_start}_{date_end}.docx"
            if date_start != date_end
            else f"exported_data/plan_{date_start}.docx"
        )
        doc.save(filename)

    def export_long_plan(self, plan):
        date_start, date_end = self.get_date_start_end(plan)
        plan_data = self.get_plan_data(plan)
        doc = Document()
        doc.add_heading(f"Recipes for {date_start} - {date_end}", 0)
        for day in plan_data["days"]:
            doc.add_heading(day["date"].strftime("%d %B %Y, %A"), level=1)
            doc.add_paragraph()
            for meal, recipe in day["meals"].items():
                doc.add_heading(f"{meal.capitalize()}", level=2)
                doc.add_paragraph(recipe["title"].upper())
                for ingredient in recipe["ingredients"]:
                    doc.add_paragraph(ingredient, style="List Bullet")
                doc.add_paragraph(recipe["cooking_instructions"])
                doc.add_paragraph()
        filename = (
            f"exported_data/recipes_{date_start}_{date_end}.docx"
            if date_start != date_end
            else f"exported_data/recipes_{date_start}.docx"
        )
        doc.save(filename)

    def export_shopping_list(self, plan):
        date_start, date_end = self.get_date_start_end(plan)
        shopping_list = self.get_shopping_list_data(plan)
        shopping_doc = Document()
        shopping_doc.add_heading("Lista zakupów", 0)
        total_list_price = 0

        for category, ingredients in shopping_list.items():
            shopping_doc.add_heading(category.capitalize(), level=1)
            for ingredient in ingredients:
                shopping_doc.add_paragraph(
                    (
                        f'{ingredient["name"]} - '
                        f'{ingredient["total_grams"]} g - '
                        f'cena: {ingredient["total_price"]:.2f} PLN'
                    ),
                    style="List Bullet",
                )
                total_list_price += ingredient["total_price"]
        shopping_doc.add_paragraph()
        shopping_doc.add_heading("KOSZT ZAKUPÓW", 1)
        shopping_doc.add_paragraph()
        shopping_doc.add_paragraph(
            f"Całkowity koszt zakupów: {total_list_price:.2f} PLN"
        )
        filename = (
            f"exported_data/shopping_list_{date_start}_{date_end}.docx"
            if date_start != date_end
            else f"exported_data/shopping_list_{date_start}.docx"
        )
        shopping_doc.save(filename)

    def do_export(self, plan):
        self.export_plan(plan)
        self.export_long_plan(plan)
        self.export_shopping_list(plan)
        print("Export done")
